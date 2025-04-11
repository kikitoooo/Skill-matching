import os
from PyPDF2 import PdfReader
from docx import Document
from django.core.files.uploadedfile import InMemoryUploadedFile

class ResumeParser:
    @staticmethod
    def parse_txt(file):
        """Парсинг txt файла"""
        if isinstance(file, InMemoryUploadedFile):
            return file.read().decode('utf-8')
        with open(file, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def parse_pdf(file):
        """Парсинг PDF файла"""
        text = ""
        if isinstance(file, InMemoryUploadedFile):
            reader = PdfReader(file)
        else:
            with open(file, 'rb') as f:
                reader = PdfReader(f)

        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    @staticmethod
    def parse_docx(file):
        """Парсинг docx файла"""
        if isinstance(file, InMemoryUploadedFile):
            doc = Document(file)
        else:
            doc = Document(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    @staticmethod
    def parse_file(file):
        """Автоматическое определение типа файла и парсинг"""
        if isinstance(file, InMemoryUploadedFile):
            filename = file.name
        else:
            if not os.path.exists(file):
                raise FileNotFoundError(f"Файл не найден: {file}")
            filename = file

        ext = os.path.splitext(filename)[1].lower()

        if ext == '.txt':
            return ResumeParser.parse_txt(file)
        elif ext == '.pdf':
            return ResumeParser.parse_pdf(file)
        elif ext == '.docx':
            return ResumeParser.parse_docx(file)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")